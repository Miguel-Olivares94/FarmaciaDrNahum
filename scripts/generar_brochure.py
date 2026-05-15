"""
Genera un brochure/manual de ventas profesional para
Sistema de Gestión Farmacéutica Dr. Nahum.
Salida: BROCHURE_FARMACIA_DR_NAHUM.docx
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).resolve().parent.parent  # raíz del proyecto
SHOTS = BASE_DIR / "screenshots"
OUT   = BASE_DIR / "docs" / "BROCHURE_FARMACIA_DR_NAHUM.docx"

# Paleta de colores — hex strings para fondos de celda, RGBColor para fuentes
NAVY_HEX    = "1A2B4A"
EMERALD_HEX = "109968"
LIGHT_HEX   = "F0F4FF"

NAVY    = RGBColor(0x1a, 0x2b, 0x4a)
EMERALD = RGBColor(0x10, 0x99, 0x68)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
GRAY    = RGBColor(0x55, 0x65, 0x81)
LIGHT   = RGBColor(0xF0, 0xF4, 0xFF)


# ─────────────────────────────── helpers ────────────────────────────────── #

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="DDDDDD"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def heading(doc, text, level=1, color=None, center=False):
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    if color:
        run.font.color.rgb = color
    return p


def body(doc, text, color=None, center=False, bold=False, size=None):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p


def add_image(doc, filename, width=Inches(5.5), caption=None):
    path = SHOTS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=width)
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.color.rgb = GRAY
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].italic = True


def section_banner(doc, title, subtitle=None):
    """Banda de color con título de sección."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, NAVY_HEX)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(14)
    if subtitle:
        p2 = cell.add_paragraph(subtitle)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.runs[0].font.color.rgb = RGBColor(0xA0, 0xD0, 0xFF)
        p2.runs[0].font.size = Pt(10)
    doc.add_paragraph()


def feature_row(doc, icon, title, description):
    """Fila con ícono, título y descripción en tabla de 2 columnas."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.columns[0].width = Cm(1.5)
    tbl.columns[1].width = Cm(14)
    left  = tbl.cell(0, 0)
    right = tbl.cell(0, 1)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_icon = left.paragraphs[0]
    p_icon.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_icon.add_run(icon)
    r.font.size = Pt(22)
    p_title = right.paragraphs[0]
    r2 = p_title.add_run(title)
    r2.bold = True
    r2.font.color.rgb = NAVY
    r2.font.size = Pt(11)
    p_desc = right.add_paragraph(description)
    p_desc.runs[0].font.color.rgb = GRAY
    p_desc.runs[0].font.size = Pt(10)
    doc.add_paragraph()


def benefit_table(doc, items):
    """Tabla de 2 columnas con beneficios."""
    rows = (len(items) + 1) // 2
    tbl  = doc.add_table(rows=rows, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    idx = 0
    for r in range(rows):
        for c in range(2):
            if idx >= len(items):
                break
            cell = tbl.cell(r, c)
            set_cell_bg(cell, LIGHT_HEX)
            set_cell_borders(cell, "CCDDFF")
            cell.width = Cm(7.5)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            run = p.add_run(items[idx])
            run.font.size = Pt(10)
            idx += 1
    doc.add_paragraph()


# ─────────────────────────────── documento ──────────────────────────────── #

def build():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── PORTADA ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, NAVY_HEX)
    cell.width = Cm(16)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    r = p.add_run("💊")
    r.font.size = Pt(48)

    p2 = cell.add_paragraph("SISTEMA DE GESTIÓN FARMACÉUTICA")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].bold = True
    p2.runs[0].font.color.rgb = RGBColor(0xA0, 0xD0, 0xFF)
    p2.runs[0].font.size = Pt(13)

    p3 = cell.add_paragraph("DR. NAHUM")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.runs[0].bold = True
    p3.runs[0].font.color.rgb = WHITE
    p3.runs[0].font.size = Pt(32)

    p4 = cell.add_paragraph("Solución digital completa para la gestión moderna de farmacias")
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.runs[0].font.color.rgb = RGBColor(0xC0, 0xD8, 0xFF)
    p4.runs[0].font.size = Pt(12)
    p4.paragraph_format.space_after = Pt(30)

    doc.add_paragraph()

    # Subtítulo portada
    p5 = doc.add_paragraph("Control total de su farmacia — desde el inventario hasta la boleta")
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.runs[0].font.color.rgb = GRAY
    p5.runs[0].italic = True
    p5.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # ── ¿QUÉ ES? ───────────────────────────────────────────────────────────
    section_banner(doc, "¿Qué es el Sistema Dr. Nahum?")

    body(doc,
         "El Sistema de Gestión Farmacéutica Dr. Nahum es una plataforma web moderna "
         "que centraliza todas las operaciones de su farmacia en un solo lugar: "
         "ventas, stock, recetas, reportes y administración de personal.",
         size=11)
    doc.add_paragraph()

    body(doc,
         "Diseñado para farmacias chilenas, cumple con los requisitos de la normativa "
         "del ISP, maneja medicamentos de venta libre y controlados, y emite boletas "
         "electrónicas. Funciona desde cualquier dispositivo con navegador — "
         "computador, tablet o smartphone.",
         size=11)
    doc.add_paragraph()

    add_image(doc, "002_dashboard_inicio.jpg",
              caption="Panel de inicio — resumen de ventas, stock y alertas en tiempo real")

    doc.add_page_break()

    # ── MÓDULOS ────────────────────────────────────────────────────────────
    section_banner(doc, "Módulos del Sistema",
                   "Todo lo que necesita para operar su farmacia")

    feature_row(doc, "🖥️", "Punto de Venta (POS)",
                "Terminal rápida para procesar ventas, buscar productos, "
                "aplicar descuentos, seleccionar método de pago y emitir "
                "boleta en segundos.")

    feature_row(doc, "💊", "Gestión de Medicamentos",
                "Inventario completo con control de lotes, fechas de vencimiento, "
                "alertas de stock bajo y búsqueda por nombre o laboratorio.")

    feature_row(doc, "📋", "Control de Recetas",
                "Validación de recetas médicas para medicamentos controlados, "
                "registro del médico, RUT del paciente y trazabilidad completa.")

    feature_row(doc, "📊", "Reportes y Análisis",
                "Informes de ventas por período, productos más vendidos, "
                "movimientos de stock y métricas para toma de decisiones.")

    feature_row(doc, "👥", "Administración de Usuarios",
                "Control de acceso por roles: Administrador, Farmacéutico y "
                "Vendedor. Cada rol ve solo lo que le corresponde.")

    feature_row(doc, "📧", "Boleta por Email",
                "El sistema envía automáticamente la boleta al cliente "
                "por correo electrónico en formato PDF.")

    doc.add_page_break()

    # ── CAPTURAS POS ───────────────────────────────────────────────────────
    section_banner(doc, "Punto de Venta en Acción")

    add_image(doc, "020_pos_nueva_venta.jpg",
              caption="Inicio de venta — búsqueda rápida de producto")
    doc.add_paragraph()

    add_image(doc, "021_pos_carrito.jpg",
              caption="Carrito de compra — cantidad, precio y total al instante")
    doc.add_paragraph()

    add_image(doc, "023_pos_metodo_pago.jpg",
              caption="Selección de método de pago: efectivo, débito o crédito")

    doc.add_page_break()

    add_image(doc, "025_pos_recibo.jpg",
              caption="Boleta generada — se puede imprimir o enviar por email")

    doc.add_page_break()

    # ── CAPTURAS MEDICAMENTOS ───────────────────────────────────────────────
    section_banner(doc, "Gestión de Medicamentos e Inventario")

    add_image(doc, "010_medicamentos_lista.jpg",
              caption="Listado completo de medicamentos con estado de stock")
    doc.add_paragraph()

    add_image(doc, "013_medicamento_detalle.jpg",
              caption="Ficha del medicamento — lotes, vencimientos y proveedores")

    doc.add_page_break()

    # ── CAPTURAS RECETAS ────────────────────────────────────────────────────
    section_banner(doc, "Control de Recetas Médicas")

    add_image(doc, "030_recetas_lista.jpg",
              caption="Registro de recetas — validación y trazabilidad completa")
    doc.add_paragraph()

    add_image(doc, "031_recetas_validar.jpg",
              caption="Validación de receta médica — datos del paciente y médico tratante")

    doc.add_page_break()

    # ── CAPTURAS REPORTES ───────────────────────────────────────────────────
    section_banner(doc, "Reportes y Análisis")

    add_image(doc, "041_reportes_ventas.jpg",
              caption="Reporte de ventas — por período, por producto o por vendedor")
    doc.add_paragraph()

    add_image(doc, "042_reportes_stock.jpg",
              caption="Reporte de stock — alertas de quiebre y productos próximos a vencer")

    doc.add_page_break()

    # ── CAPTURAS ADMIN ──────────────────────────────────────────────────────
    section_banner(doc, "Panel de Administración")

    add_image(doc, "051_admin_usuarios.jpg",
              caption="Gestión de usuarios y roles — control total de accesos")
    doc.add_paragraph()

    add_image(doc, "056_admin_auditoría.jpg",
              caption="Registro de auditoría — historial completo de acciones en el sistema")

    doc.add_page_break()

    # ── BENEFICIOS ──────────────────────────────────────────────────────────
    section_banner(doc, "¿Por qué elegir este sistema?",
                   "Beneficios concretos para su farmacia")

    benefit_table(doc, [
        "✅  Reduce errores en ventas y stock",
        "✅  Ahorra tiempo en cada transacción",
        "✅  Acceso desde cualquier dispositivo",
        "✅  Sin instalaciones complicadas",
        "✅  Boleta electrónica automática",
        "✅  Control de medicamentos controlados",
        "✅  Alertas de vencimiento y stock bajo",
        "✅  Respaldo en la nube — sus datos seguros",
        "✅  Reportes para tomar mejores decisiones",
        "✅  Soporte y capacitación incluidos",
    ])

    doc.add_page_break()

    # ── CONTACTO ────────────────────────────────────────────────────────────
    section_banner(doc, "Solicite una Demostración")

    doc.add_paragraph()
    body(doc, "¿Le interesa implementar este sistema en su farmacia?",
         bold=True, center=True, size=13, color=NAVY)
    body(doc, "Agendamos una demostración sin compromiso.",
         center=True, size=11, color=GRAY)
    doc.add_paragraph()

    tbl2 = doc.add_table(rows=4, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    contactos = [
        ("📱  WhatsApp",    "+56 9 4265 2487"),
        ("📧  Email",       "admin@farmaciadrnahum.cl"),
        ("🌐  Demo en línea", "https://web-production-b4da9.up.railway.app"),
        ("⏰  Atención",    "Lunes a Viernes, 09:00 – 18:00 hrs"),
    ]
    for i, (label, value) in enumerate(contactos):
        c0 = tbl2.cell(i, 0)
        c1 = tbl2.cell(i, 1)
        set_cell_bg(c0, NAVY_HEX)
        set_cell_bg(c1, LIGHT_HEX)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(6)
        p0.paragraph_format.space_after  = Pt(6)
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.color.rgb = WHITE
        r0.font.size = Pt(10)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(6)
        p1.paragraph_format.space_after  = Pt(6)
        r1 = p1.add_run(value)
        r1.font.color.rgb = NAVY
        r1.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    # Pie de página
    footer = doc.add_paragraph(
        "Sistema de Gestión Farmacéutica Dr. Nahum  —  Desarrollado en Chile  —  2026"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.color.rgb = GRAY
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].italic = True

    doc.save(str(OUT))
    print(f"Brochure generado: {OUT}")


if __name__ == "__main__":
    build()
