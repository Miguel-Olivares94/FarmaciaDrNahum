#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generar_manual.py
=================
Script único para generar el Manual Completo del Sistema Farmacia Dr Nahúm.
Genera un archivo Word (.docx) profesional con toda la documentación.

Uso:
    python generar_manual.py

Requiere:
    pip install python-docx
"""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────
# CONFIGURACIÓN
# ─────────────────────────────────────────────────────────────────────
BASE_DIR   = Path(__file__).parent.parent  # raíz del proyecto
OUTPUT     = BASE_DIR / "docs" / "MANUAL_COMPLETO_FARMACIA.docx"
SHOTS_DIR  = BASE_DIR / "screenshots"
URL_SISTEMA = "http://localhost:5000"

# Paleta de colores (Navy + Emerald)
C_NAVY    = RGBColor(15,  41,  66)   # #0f2942
C_PRIMARY = RGBColor(29, 111, 164)   # #1d6fa4
C_GREEN   = RGBColor( 5, 150, 105)   # #059669
C_AMBER   = RGBColor(217, 119,   6)  # #d97706
C_RED     = RGBColor(220,  38,  38)  # #dc2626
C_GRAY    = RGBColor(100, 116, 139)  # #64748b
C_WHITE   = RGBColor(255, 255, 255)
C_LIGHT   = RGBColor(241, 245, 249)  # #f1f5f9

FONT = "Calibri"


# ─────────────────────────────────────────────────────────────────────
# HELPERS DE FORMATO
# ─────────────────────────────────────────────────────────────────────

def _run(para, text, bold=False, italic=False, size=11, color=None, font=FONT):
    run = para.add_run(text)
    run.font.name  = font
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def _cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _cell_text(cell, text, bold=False, color=None, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    _run(p, text, bold=bold, color=color, size=size)


def add_page_title(doc, number, title):
    """Título de capítulo: número grande + texto."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    _run(p, f"{number}. ", bold=True, size=22, color=C_PRIMARY)
    _run(p, title,         bold=True, size=22, color=C_NAVY)
    # Línea divisora
    border_p = doc.add_paragraph()
    border_p.paragraph_format.space_before = Pt(0)
    border_p.paragraph_format.space_after  = Pt(10)
    pPr = border_p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1d6fa4')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section(doc, title):
    """Subtítulo de sección."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    _run(p, title, bold=True, size=13, color=C_PRIMARY)


def add_body(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    _run(p, text, size=11)
    return p


def add_bullet(doc, text, sub=False):
    style = 'List Bullet 2' if sub else 'List Bullet'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    _run(p, text, size=10.5)


def add_numbered(doc, text, sub=False):
    style = 'List Number 2' if sub else 'List Number'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    _run(p, text, size=10.5)


def add_step(doc, n, title, desc):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(5)
    _run(p, f"{title}: ", bold=True, size=11, color=C_PRIMARY)
    _run(p, desc, size=11)


def add_note(doc, kind, text):
    """Caja de nota/advertencia/alerta."""
    colors = {
        'info':    ('eff6ff', C_PRIMARY, 'ℹ️  Nota'),
        'warning': ('fffbeb', C_AMBER,   '⚠️  Atención'),
        'danger':  ('fff5f5', C_RED,     '🚫  Importante'),
        'tip':     ('f0fdf4', C_GREEN,   '✅  Consejo'),
    }
    bg, color, label = colors.get(kind, colors['info'])
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(0.5)
    p.paragraph_format.right_indent  = Cm(0.5)
    p.paragraph_format.space_before  = Pt(6)
    p.paragraph_format.space_after   = Pt(8)
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  bg)
    pPr.append(shd)
    _run(p, f"{label}: ", bold=True, size=10.5, color=color)
    _run(p, text, size=10.5)


def add_screenshot(doc, filename, caption=""):
    """Inserta imagen si existe, o placeholder si no."""
    path = SHOTS_DIR / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(5.5))
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, "", italic=True, size=9, color=C_GRAY)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, f"[Captura: {filename}]", italic=True, size=9, color=C_GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table_header_row(table, headers, bg="0f2942"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        _cell_bg(cell, bg)
        _cell_text(cell, h, bold=True, color=C_WHITE, size=10,
                   align=WD_ALIGN_PARAGRAPH.CENTER)


# ─────────────────────────────────────────────────────────────────────
# PORTADA
# ─────────────────────────────────────────────────────────────────────

def build_cover(doc):
    for _ in range(4):
        doc.add_paragraph()

    # Logo textual
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "💊", size=48)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    _run(p, "FARMACIA DR NAHÚM", bold=True, size=28, color=C_NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Sistema de Gestión Farmacéutica", italic=True, size=16, color=C_PRIMARY)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "MANUAL COMPLETO DEL USUARIO", bold=True, size=20, color=C_NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Guía para Vendedores, Supervisores y Administradores",
         italic=True, size=12, color=C_GRAY)

    for _ in range(3):
        doc.add_paragraph()

    # Tabla de metadatos
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    meta = [
        ("Versión",     "1.0"),
        ("Fecha",       datetime.now().strftime("%d de %B de %Y")),
        ("Framework",   "Django 6 + Bootstrap 5"),
        ("Puerto",      "http://localhost:5000"),
    ]
    for i, (k, v) in enumerate(meta):
        row = table.rows[i]
        _cell_bg(row.cells[0], "1d6fa4")
        _cell_text(row.cells[0], k, bold=True, color=C_WHITE, size=10)
        _cell_text(row.cells[1], v, size=10)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────
# ÍNDICE
# ─────────────────────────────────────────────────────────────────────

def build_toc(doc):
    add_page_title(doc, "Índice", "de Contenidos")
    toc = [
        ("1", "Introducción al Sistema"),
        ("2", "Requisitos y Acceso"),
        ("   2.1", "Iniciar Sesión"),
        ("   2.2", "Roles y Permisos"),
        ("3", "Panel Principal (Home)"),
        ("4", "Gestión de Medicamentos"),
        ("   4.1", "Ver y Filtrar Medicamentos"),
        ("   4.2", "Agregar Medicamento"),
        ("   4.3", "Editar y Eliminar"),
        ("   4.4", "Control de Stock"),
        ("5", "Terminal POS (Punto de Venta)"),
        ("   5.1", "Crear Nueva Venta"),
        ("   5.2", "Agregar Productos"),
        ("   5.3", "Validación de Recetas"),
        ("   5.4", "Procesar Pago"),
        ("   5.5", "Anular Venta"),
        ("6", "Control de Lotes e Inventario"),
        ("   6.1", "Gestor de Lotes"),
        ("   6.2", "Dashboard Inventario"),
        ("   6.3", "Reporte de Lotes"),
        ("7", "Gestión de Proveedores"),
        ("8", "Reportes y Dashboard"),
        ("9", "Administración del Sistema"),
        ("   9.1", "Gestión de Usuarios"),
        ("   9.2", "Auditoría"),
        ("10", "Referencia Rápida y Solución de Problemas"),
    ]
    for num, title in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        indent = num.startswith("   ")
        if indent:
            p.paragraph_format.left_indent = Cm(1)
            _run(p, num.strip(), size=10, color=C_PRIMARY)
            _run(p, f"  {title}", size=10)
        else:
            _run(p, f"{num}. ", bold=True, size=11, color=C_PRIMARY)
            _run(p, title, bold=True, size=11)
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────
# CAP 1 — INTRODUCCIÓN
# ─────────────────────────────────────────────────────────────────────

def build_cap1(doc):
    add_page_title(doc, "1", "Introducción al Sistema")

    add_body(doc,
        "El Sistema de Gestión Farmacéutica Farmacia Dr Nahúm es una aplicación web "
        "desarrollada con Django 6 y Bootstrap 5 que permite administrar de forma integral "
        "todas las operaciones de una farmacia moderna.")

    add_section(doc, "¿Qué puedo hacer con el sistema?")
    bullets = [
        "Controlar el inventario de medicamentos en tiempo real",
        "Registrar ventas mediante Terminal POS integrada",
        "Validar recetas médicas (simple, retenida, controlada)",
        "Gestionar lotes con control de vencimientos (FIFO)",
        "Administrar proveedores y órdenes de compra",
        "Generar reportes de ventas, stock y finanzas",
        "Controlar accesos con sistema de roles y permisos",
        "Auditar todas las operaciones del sistema",
    ]
    for b in bullets:
        add_bullet(doc, b)

    add_section(doc, "Tecnología")
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    data = [
        ("Backend",    "Python 3.14 + Django 6"),
        ("Frontend",   "Bootstrap 5.3 + Bootstrap Icons"),
        ("Base de Datos", "SQLite (desarrollo) / MySQL (producción)"),
        ("Puerto",     "http://localhost:5000"),
        ("Acceso",     "Navegador web: Chrome, Firefox, Safari, Edge"),
    ]
    for i, (k, v) in enumerate(data):
        _cell_bg(table.rows[i].cells[0], "f1f5f9")
        _cell_text(table.rows[i].cells[0], k, bold=True, size=10)
        _cell_text(table.rows[i].cells[1], v, size=10)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────
# CAP 2 — ACCESO
# ─────────────────────────────────────────────────────────────────────

def build_cap2(doc):
    add_page_title(doc, "2", "Requisitos y Acceso")

    add_section(doc, "2.1  Iniciar Sesión")
    add_body(doc, "Para acceder al sistema:")
    add_step(doc, 1, "Abrir navegador", "Chrome, Firefox, Safari o Edge")
    add_step(doc, 2, "Ir a la URL", URL_SISTEMA)
    add_step(doc, 3, "Ingresar usuario", "Tu nombre de usuario (ej: admin, vendedor1)")
    add_step(doc, 4, "Ingresar contraseña", "La contraseña asignada por el administrador")
    add_step(doc, 5, "Clic en 'Iniciar Sesión'", "El sistema carga tu panel según tu rol")

    add_note(doc, 'warning',
        "Nunca compartas tu contraseña. Si la olvidaste, pide al administrador que la restablezca.")

    add_screenshot(doc, "login.jpg", "Pantalla de inicio de sesión")

    add_section(doc, "2.2  Roles y Permisos")
    add_body(doc, "El sistema tiene 4 niveles de acceso:")

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    add_table_header_row(table, ["Rol", "Acceso", "Restricciones"])
    roles = [
        ("Admin / Superusuario", "Acceso total: usuarios, reportes, configuración, anulaciones",
         "Ninguna"),
        ("Gerente",  "Inventario, reportes globales, lotes, proveedores",
         "No puede anular ventas de otros"),
        ("Vendedor", "Terminal POS, ver medicamentos, historial propio",
         "Sin acceso a reportes financieros"),
        ("Contador", "Reportes financieros, auditoría",
         "Solo lectura, sin ventas"),
    ]
    for i, (rol, acceso, rest) in enumerate(roles, 1):
        row = table.rows[i]
        _cell_bg(row.cells[0], "dbeafe")
        _cell_text(row.cells[0], rol,    bold=True, size=10)
        _cell_text(row.cells[1], acceso, size=10)
        _cell_text(row.cells[2], rest,   size=10, color=C_GRAY)

    add_note(doc, 'info',
        "Los usuarios disponibles por defecto son: admin, vendedor1, vendedor2, gerente, contador.")

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────
# CAP 3 — PANEL PRINCIPAL
# ─────────────────────────────────────────────────────────────────────

def build_cap3(doc):
    add_page_title(doc, "3", "Panel Principal (Home)")

    add_body(doc,
        "Al iniciar sesión llegas al Panel Principal. Desde aquí puedes ver "
        "el resumen del sistema y navegar a cualquier módulo.")

    add_section(doc, "Elementos del Panel")
    items = [
        ("Hero Card",        "Muestra el nombre del sistema, estadísticas clave y botones de acceso rápido"),
        ("KPIs",             "Tarjetas con: total medicamentos, agotados, en stock, usuario activo"),
        ("Accesos Rápidos",  "Botones directos a: Terminal POS, Medicamentos, Historial Ventas, Proveedores, Dashboard Inventario"),
        ("Información",      "Versión del sistema, rol del usuario, links de redes sociales"),
        ("Navbar superior",  "Acceso a todos los módulos desde el menú superior"),
        ("Sidebar (desktop)","Panel izquierdo con navegación completa en pantallas grandes"),
        ("Barra inferior (móvil)", "Navegación de 5 accesos en teléfonos: Inicio, Stock, POS, Ventas, Panel"),
    ]
    table = doc.add_table(rows=len(items)+1, cols=2)
    table.style = 'Table Grid'
    add_table_header_row(table, ["Elemento", "Descripción"])
    for i, (elem, desc) in enumerate(items, 1):
        _cell_bg(table.rows[i].cells[0], "f8fafc")
        _cell_text(table.rows[i].cells[0], elem, bold=True, size=10)
        _cell_text(table.rows[i].cells[1], desc, size=10)

    add_screenshot(doc, "home.jpg", "Panel principal del sistema")
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────
# CAP 4 — MEDICAMENTOS
# ─────────────────────────────────────────────────────────────────────

def build_cap4(doc):
    add_page_title(doc, "4", "Gestión de Medicamentos")

    add_section(doc, "4.1  Ver y Filtrar Medicamentos")
    add_body(doc, "Menú → Inventario → Medicamentos")
    add_step(doc, 1, "Tabla principal",    "Muestra todos los medicamentos: SKU, nombre, laboratorio, stock, precio, tipo")
    add_step(doc, 2, "Filtros superiores", "Botones para filtrar por tipo de venta: Todos, Con Receta, Receta Simple, Receta Retenida, Controlados, Venta Libre")
    add_step(doc, 3, "KPIs",              "Tarjetas con: total, stock total, bajo stock, agotados")
    add_step(doc, 4, "Colores de stock",   "🟢 Alto (>700 unid.) · 🟡 Medio (100-700) · 🔴 Bajo (<100)")

    add_screenshot(doc, "medicamentos_lista.jpg", "Listado de medicamentos con filtros")

    add_section(doc, "Tipos de Venta")
    tipos = [
        ("Venta Libre",       "🟢", "Sin receta. Ej: paracetamol, ibuprofeno genérico"),
        ("Receta Simple",     "🔵", "Requiere receta, no se retiene. Ej: antibióticos"),
        ("Receta Retenida",   "🟡", "Requiere receta y se archiva en la farmacia"),
        ("Controlado",        "🔴", "Estupefaciente/psicotrópico, requiere aprobación de supervisor"),
    ]
    table = doc.add_table(rows=len(tipos)+1, cols=3)
    table.style = 'Table Grid'
    add_table_header_row(table, ["Tipo", "Indicador", "Descripción"])
    for i, (tipo, ind, desc) in enumerate(tipos, 1):
        _cell_text(table.rows[i].cells[0], tipo, bold=True, size=10)
        _cell_text(table.rows[i].cells[1], ind,  size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(table.rows[i].cells[2], desc, size=10)

    add_section(doc, "4.2  Agregar Medicamento")
    add_body(doc, "Menú → Inventario → Agregar (o botón 'Agregar Medicamento')")
    add_step(doc, 1, "SKU",               "Código único del medicamento (ej: MED-001)")
    add_step(doc, 2, "Nombre",            "Nombre completo del medicamento")
    add_step(doc, 3, "Principio activo",  "Componente activo principal")
    add_step(doc, 4, "Laboratorio",       "Fabricante o distribuidor")
    add_step(doc, 5, "Tipo de venta",     "Libre / Receta Simple / Receta Retenida / Controlado")
    add_step(doc, 6, "Stock y precio",    "Cantidad inicial disponible y precio de venta")
    add_step(doc, 7, "Guardar",           "Clic en 'Guardar' — el medicamento aparece en el catálogo")

    add_note(doc, 'warning', "El SKU debe ser único. Si ya existe, el sistema rechazará el formulario.")

    add_section(doc, "4.3  Editar y Eliminar")
    add_bullet(doc, "Editar: Botón ✏️ en la fila del medicamento → modifica datos → Guardar")
    add_bullet(doc, "Eliminar: Botón 🗑️ → confirmar eliminación. Solo admin puede eliminar")
    add_note(doc, 'danger', "Eliminar un medicamento borra su historial. Prefiere desactivarlo cambiando el stock a 0.")

    add_section(doc, "4.4  Control de Stock")
    add_bullet(doc, "Stock alto (>700): indicador verde")
    add_bullet(doc, "Stock medio (100–700): indicador amarillo")
    add_bullet(doc, "Stock bajo (<100): indicador rojo — notificar al supervisor")
    add_bullet(doc, "Stock cero: medicamento bloqueado para venta")
    add_note(doc, 'tip', "El sistema descuenta el stock automáticamente al confirmar cada venta.")

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────
# CAP 5 — POS
# ─────────────────────────────────────────────────────────────────────

def build_cap5(doc):
    add_page_title(doc, "5", "Terminal POS — Punto de Venta")

    add_body(doc,
        "La Terminal POS es el módulo central para registrar ventas. "
        "Acceso: Menú → Ventas → Terminal POS  (o botón azul en barra inferior en móvil).")

    add_section(doc, "5.1  Crear Nueva Venta")
    add_step(doc, 1, "Abrir Terminal POS", "Menú superior o botón 'POS' en barra de navegación")
    add_step(doc, 2, "Carrito vacío",      "Se muestra el área de venta lista para agregar productos")
    add_step(doc, 3, "Seleccionar cliente","Opcional: buscar cliente registrado o dejar anónimo")

    add_screenshot(doc, "pos_terminal.jpg", "Terminal POS - Pantalla principal")

    add_section(doc, "5.2  Agregar Productos")
    add_body(doc, "Tres formas de agregar medicamentos:")

    add_bullet(doc, "Búsqueda por nombre: Escribe el nombre en el buscador → selecciona de la lista")
    add_bullet(doc, "Código de barras: Escanea con lector físico → se agrega automáticamente")
    add_bullet(doc, "Catálogo: Haz clic en el medicamento de la grilla de productos")

    add_body(doc, "Al agregar un medicamento:", space_after=4)
    add_numbered(doc, "Especifica la cantidad deseada")
    add_numbered(doc, "El sistema valida: stock disponible, vencimiento del lote, requisito de receta")
    add_numbered(doc, "Se calcula el subtotal (cantidad × precio)")
    add_numbered(doc, "Puedes modificar la cantidad o eliminar el ítem del carrito")

    add_section(doc, "5.3  Validación de Recetas")
    add_note(doc, 'warning',
        "Al agregar un medicamento que requiere receta, el sistema solicita los datos antes de continuar.")

    tipos_receta = [
        ("Venta Libre",     "Sin requisitos — venta inmediata"),
        ("Receta Simple",   "Ingresa: RUT médico, nombre médico, RUT paciente, fecha receta"),
        ("Receta Retenida", "Igual a receta simple + farmacia archiva copia física de la receta"),
        ("Controlado",      "Requiere autorización del supervisor antes de proceder"),
    ]
    for tipo, req in tipos_receta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        _run(p, f"• {tipo}: ", bold=True, size=10.5, color=C_PRIMARY)
        _run(p, req, size=10.5)

    add_section(doc, "5.4  Procesar Pago")
    add_step(doc, 1, "Revisar carrito",     "Verifica ítems, cantidades y precios")
    add_step(doc, 2, "Seleccionar medio",   "Efectivo / Tarjeta débito / Tarjeta crédito / Transferencia")
    add_step(doc, 3, "Ingresar monto",      "Cuánto dinero recibiste del cliente")
    add_step(doc, 4, "Cambio",              "El sistema calcula el vuelto automáticamente")
    add_step(doc, 5, "Confirmar",           "Clic en 'Completar Venta' — se genera la boleta")
    add_step(doc, 6, "Boleta",              "Se muestra boleta en pantalla, disponible para imprimir o enviar por email")

    add_note(doc, 'danger',
        "Una venta confirmada no puede editarse. Para corregir errores usa la opción 'Anular'.")

    add_screenshot(doc, "pos_pago.jpg", "Pantalla de procesamiento de pago")

    add_section(doc, "5.5  Anular Venta")
    add_body(doc, "Solo administradores pueden anular ventas:")
    add_step(doc, 1, "Ir a Historial de Ventas", "Menú → Ventas → Historial de Ventas")
    add_step(doc, 2, "Buscar la venta",           "Por número, fecha o cliente")
    add_step(doc, 3, "Clic en 'Anular'",          "Ingresar motivo de anulación")
    add_step(doc, 4, "Confirmar",                  "El stock se repone automáticamente al anular")

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────
# CAP 6 — LOTES E INVENTARIO
# ─────────────────────────────────────────────────────────────────────

def build_cap6(doc):
    add_page_title(doc, "6", "Control de Lotes e Inventario")

    add_body(doc,
        "El sistema usa método FIFO (primero en entrar, primero en salir) "
        "para gestionar los lotes de medicamentos y controlar vencimientos.")

    add_section(doc, "6.1  Gestor de Lotes")
    add_body(doc, "Menú → Inventario → Gestor de Lotes")
    add_bullet(doc, "Ver todos los lotes próximos a vencer (< 30 días)")
    add_bullet(doc, "Filtrar lotes por medicamento específico")
    add_bullet(doc, "Ver cantidad disponible por lote y fecha de vencimiento")

    add_section(doc, "6.2  Dashboard de Inventario")
    add_body(doc, "Menú → Inventario → Dashboard Inventario")
    add_bullet(doc, "Medicamentos con alerta de vencimiento (< 7 días)")
    add_bullet(doc, "Lotes vencidos con stock disponible (deben retirarse)")
    add_bullet(doc, "Lotes críticos: vencen en menos de 7 días")
    add_bullet(doc, "Estadísticas: total medicamentos, agotados")

    add_note(doc, 'danger',
        "Los medicamentos vencidos aparecen marcados en rojo. NUNCA vender un medicamento vencido.")

    add_section(doc, "6.3  Reporte de Lotes")
    add_body(doc, "Menú → Ventas → Reportes")
    add_bullet(doc, "Lotes vigentes: con stock y sin vencer")
    add_bullet(doc, "Lotes críticos: vencen en < 7 días")
    add_bullet(doc, "Lotes vencidos: deben ser retirados del inventario")
    add_bullet(doc, "Totales de stock por estado")

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────
# CAP 7 — PROVEEDORES
# ─────────────────────────────────────────────────────────────────────

def build_cap7(doc):
    add_page_title(doc, "7", "Gestión de Proveedores")

    add_body(doc, "Menú → Inventario → Proveedores")

    add_section(doc, "Ver Proveedores")
    add_bullet(doc, "Lista con todos los proveedores registrados")
    add_bullet(doc, "Ver detalles: contacto, teléfono, email, dirección")

    add_section(doc, "Agregar Proveedor")
    add_step(doc, 1, "Nuevo Proveedor", "Clic en botón 'Nuevo Proveedor'")
    add_step(doc, 2, "Nombre",          "Razón social del proveedor")
    add_step(doc, 3, "Contacto",        "Nombre del representante o contacto")
    add_step(doc, 4, "Teléfono",        "Número de contacto")
    add_step(doc, 5, "Email",           "Correo electrónico de contacto")
    add_step(doc, 6, "Guardar",         "El proveedor queda registrado en el sistema")

    add_section(doc, "Editar / Eliminar")
    add_bullet(doc, "Editar: botón ✏️ → modificar datos → Guardar")
    add_bullet(doc, "Eliminar: botón 🗑️ → confirmar. Solo si no tiene medicamentos asociados")

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────
# CAP 8 — REPORTES
# ─────────────────────────────────────────────────────────────────────

def build_cap8(doc):
    add_page_title(doc, "8", "Reportes y Dashboard")

    add_section(doc, "Dashboard Principal")
    add_body(doc, "Menú → Dashboard")
    add_bullet(doc, "KPIs: medicamentos activos, total vendido, transacciones, unidades vendidas")
    add_bullet(doc, "Tabla de medicamentos con bajo stock y nivel de alerta")
    add_bullet(doc, "Estadísticas de ventas: semanales, mensuales y anuales")
    add_bullet(doc, "Accesos rápidos: Terminal POS, Inventario, Historial, Dashboard Lotes")

    add_section(doc, "Dashboard de Inventario")
    add_body(doc, "Menú → Inventario → Dashboard Inventario")
    add_bullet(doc, "Medicamentos próximos a vencer")
    add_bullet(doc, "Lotes vencidos con stock")
    add_bullet(doc, "Alertas críticas de stock")

    add_section(doc, "Historial de Ventas")
    add_body(doc, "Menú → Ventas → Historial Ventas")
    add_bullet(doc, "Lista de todas las ventas con fecha, hora, total y estado")
    add_bullet(doc, "Ver detalle de cada venta (medicamentos, cantidades, pagos)")
    add_bullet(doc, "Imprimir boleta de ventas anteriores")
    add_bullet(doc, "Anular ventas (solo administradores)")

    add_section(doc, "Reporte de Lotes")
    add_body(doc, "Menú → Ventas → Reportes")
    add_bullet(doc, "Stock vigente, crítico y vencido con totales")
    add_bullet(doc, "Ordenado por fecha de vencimiento")

    add_note(doc, 'info',
        "Los reportes están disponibles según el rol del usuario. "
        "Vendedores solo ven su historial propio.")

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────
# CAP 9 — ADMINISTRACIÓN
# ─────────────────────────────────────────────────────────────────────

def build_cap9(doc):
    add_page_title(doc, "9", "Administración del Sistema")
    add_note(doc, 'warning', "Esta sección es solo para usuarios con rol Administrador.")

    add_section(doc, "9.1  Gestión de Usuarios")
    add_body(doc, "Los usuarios se gestionan desde el panel de administración de Django:")
    add_step(doc, 1, "Ir a /admin",       "Accede a http://localhost:5000/admin con usuario admin")
    add_step(doc, 2, "Sección Usuarios",  "Clic en 'Usuarios' bajo la sección 'Autenticación'")
    add_step(doc, 3, "Nuevo usuario",     "Botón '+ Agregar Usuario'")
    add_step(doc, 4, "Datos",             "Username, contraseña, nombre, email, rol (is_staff / is_superuser)")
    add_step(doc, 5, "Guardar",           "El usuario puede iniciar sesión inmediatamente")

    add_body(doc, "Restablecer contraseña desde terminal:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1e293b')
    pPr.append(shd)
    _run(p, "python manage.py changepassword <usuario>", size=10, color=C_WHITE,
         font="Courier New")

    add_section(doc, "9.2  Auditoría")
    add_bullet(doc, "El sistema registra automáticamente todas las operaciones críticas")
    add_bullet(doc, "Accesos, ventas, modificaciones y errores quedan en el log de auditoría")
    add_bullet(doc, "Los logs se guardan en la carpeta /logs/ del proyecto")
    add_bullet(doc, "Solo administradores pueden ver el log completo de auditoría")

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────
# CAP 10 — REFERENCIA RÁPIDA
# ─────────────────────────────────────────────────────────────────────

def build_cap10(doc):
    add_page_title(doc, "10", "Referencia Rápida y Solución de Problemas")

    add_section(doc, "Flujo Operativo Diario")
    flujo = [
        ("08:00", "Revisar Dashboard → alertas de stock bajo y vencimientos"),
        ("08:15", "Abrir Terminal POS → verificar que funcione correctamente"),
        ("Durante el día", "Atender clientes → registrar ventas en Terminal POS"),
        ("Al recibir receta", "Validar receta en sistema antes de dispensar"),
        ("Medicamento controlado", "Llamar a supervisor para autorización"),
        ("Cierre del día", "Revisar historial de ventas del día"),
        ("Al cerrar", "Guardar recetas retenidas en archivo cronológico"),
    ]
    table = doc.add_table(rows=len(flujo)+1, cols=2)
    table.style = 'Table Grid'
    add_table_header_row(table, ["Momento", "Acción"])
    for i, (hora, accion) in enumerate(flujo, 1):
        _cell_bg(table.rows[i].cells[0], "eff6ff")
        _cell_text(table.rows[i].cells[0], hora,   bold=True, size=10)
        _cell_text(table.rows[i].cells[1], accion, size=10)

    add_section(doc, "Errores Comunes y Soluciones")
    errores = [
        ("Stock insuficiente",
         "El medicamento no tiene la cantidad solicitada. Verifica el stock actual."),
        ("Requiere receta",
         "El medicamento exige receta médica. Solicita al cliente que la presente."),
        ("Medicamento vencido",
         "La fecha de vencimiento ya pasó. NO VENDER. Notifica al supervisor inmediatamente."),
        ("Usuario no autorizado",
         "Tu rol no tiene permiso para esta acción. Contacta al administrador."),
        ("Error al guardar",
         "Verifica que todos los campos obligatorios estén completos y el SKU sea único."),
        ("La página no carga",
         "Verifica que el servidor esté activo: python manage.py runserver 5000"),
    ]
    for error, sol in errores:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        _run(p, f"❌ {error}: ", bold=True, size=10.5, color=C_RED)
        _run(p, sol, size=10.5)

    add_section(doc, "Usuarios del Sistema por Defecto")
    table2 = doc.add_table(rows=7, cols=3)
    table2.style = 'Table Grid'
    add_table_header_row(table2, ["Usuario", "Contraseña", "Rol"])
    usuarios = [
        ("admin",     "admin1234",  "Superusuario (acceso total)"),
        ("vendedor1", "ver en admin", "Vendedor"),
        ("vendedor2", "ver en admin", "Vendedor"),
        ("gerente",   "ver en admin", "Gerente / Staff"),
        ("contador",  "ver en admin", "Contador / Staff"),
        ("dr_nahum",  "ver en admin", "Usuario base"),
    ]
    for i, (user, pwd, rol) in enumerate(usuarios, 1):
        _cell_text(table2.rows[i].cells[0], user, bold=True, size=10)
        _cell_text(table2.rows[i].cells[1], pwd,  size=10, color=C_PRIMARY)
        _cell_text(table2.rows[i].cells[2], rol,  size=10)

    add_note(doc, 'danger',
        "Cambia las contraseñas por defecto antes de usar el sistema en producción.")

    add_section(doc, "Soporte y Contacto")
    add_bullet(doc, "📧 Email: admin@farmaciadrnahum.cl")
    add_bullet(doc, "📱 WhatsApp: +56 9 4265 2487")
    add_bullet(doc, f"🌐 Sistema: {URL_SISTEMA}")
    add_bullet(doc, "⏰ Horario soporte: Lunes a Viernes, 09:00 – 18:00 hrs")

    # Pie de página final
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "© 2026 Farmacia Dr Nahúm · Manual del Usuario v1.0 · ",
         size=9, color=C_GRAY)
    _run(p, datetime.now().strftime("%d/%m/%Y"),
         size=9, color=C_GRAY)


# ─────────────────────────────────────────────────────────────────────
# FUNCIÓN PRINCIPAL
# ─────────────────────────────────────────────────────────────────────

def generar_manual():
    print("=" * 60)
    print("  GENERANDO MANUAL FARMACIA DR NAHÚM")
    print("=" * 60)

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # Estilo base
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(11)

    print("  Construyendo portada...")
    build_cover(doc)
    print("  Construyendo índice...")
    build_toc(doc)
    print("  Cap. 1 - Introducción...")
    build_cap1(doc)
    print("  Cap. 2 - Acceso y roles...")
    build_cap2(doc)
    print("  Cap. 3 - Panel principal...")
    build_cap3(doc)
    print("  Cap. 4 - Medicamentos...")
    build_cap4(doc)
    print("  Cap. 5 - Terminal POS...")
    build_cap5(doc)
    print("  Cap. 6 - Lotes e inventario...")
    build_cap6(doc)
    print("  Cap. 7 - Proveedores...")
    build_cap7(doc)
    print("  Cap. 8 - Reportes...")
    build_cap8(doc)
    print("  Cap. 9 - Administración...")
    build_cap9(doc)
    print("  Cap. 10 - Referencia rápida...")
    build_cap10(doc)

    doc.save(str(OUTPUT))

    print()
    print("=" * 60)
    print(f"  ✅ Manual generado correctamente")
    print(f"  📄 Archivo: {OUTPUT.name}")
    print(f"  📁 Ubicación: {OUTPUT.parent}")
    print(f"  📊 Capítulos: 10")
    print("=" * 60)
    print()
    print("  Para agregar capturas de pantalla:")
    print(f"  Coloca imágenes .jpg en: {SHOTS_DIR}")
    print("  Nombres esperados: login.jpg, home.jpg,")
    print("  medicamentos_lista.jpg, pos_terminal.jpg, pos_pago.jpg")
    print("=" * 60)


if __name__ == "__main__":
    generar_manual()
